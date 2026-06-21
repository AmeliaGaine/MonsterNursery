#!/usr/bin/env python3
"""Compile all 49 episode scene-beat files into a single .docx (manuscript.docx)."""
import os, re, glob
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCENES_DIR = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(SCENES_DIR, "..", "..", "..", "manuscript.docx")
OUT = os.path.abspath(OUT)

INLINE_RE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)')

def add_inline(paragraph, text):
    """Add text to a paragraph, honoring **bold**, *italic*, `code`."""
    for chunk in INLINE_RE.split(text):
        if not chunk:
            continue
        if chunk.startswith('**') and chunk.endswith('**'):
            r = paragraph.add_run(chunk[2:-2]); r.bold = True
        elif chunk.startswith('*') and chunk.endswith('*'):
            r = paragraph.add_run(chunk[1:-1]); r.italic = True
        elif chunk.startswith('`') and chunk.endswith('`'):
            r = paragraph.add_run(chunk[1:-1]); r.font.name = 'Consolas'
        else:
            paragraph.add_run(chunk)

def render_file(doc, path):
    with open(path, encoding='utf-8') as f:
        lines = f.read().splitlines()

    in_html_comment = False
    for raw in lines:
        line = raw.rstrip()

        # skip HTML comments (the MICRO: directive line)
        if line.strip().startswith('<!--'):
            in_html_comment = not line.strip().endswith('-->')
            continue
        if in_html_comment:
            if line.strip().endswith('-->'):
                in_html_comment = False
            continue

        if not line.strip():
            continue

        # horizontal rule
        if re.match(r'^-{3,}$', line.strip()):
            doc.add_paragraph().add_run('— · —').italic = True
            continue

        # headings
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            # episode title (# Ep N) -> Heading 1 on a fresh page
            if level == 1:
                p = doc.add_heading(level=1)
                add_inline(p, text)
            elif level == 2:
                p = doc.add_heading(level=2)
                add_inline(p, text)
            else:
                p = doc.add_heading(level=3)
                add_inline(p, text)
            continue

        # blockquote
        if line.lstrip().startswith('>'):
            text = line.lstrip()[1:].lstrip()
            p = doc.add_paragraph(style='Intense Quote')
            add_inline(p, text)
            continue

        # list items (handle nesting by leading spaces)
        lm = re.match(r'^(\s*)([-*])\s+(.*)$', line)
        if lm:
            indent = len(lm.group(1))
            text = lm.group(3)
            style = 'List Bullet' if indent < 2 else 'List Bullet 2'
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style='List Bullet')
            add_inline(p, text)
            continue

        # normal paragraph
        p = doc.add_paragraph()
        add_inline(p, line)

def main():
    doc = Document()
    # base font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # Title page
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('THE MONSTER NURSERY — BOOK 1')
    r.bold = True; r.font.size = Pt(24)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run('Scene-Beats Outline — All 49 Episodes')
    rs.font.size = Pt(14); rs.italic = True
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Cozy LitRPG · compiled from the Outline Forge build').italic = True
    doc.add_page_break()

    files = sorted(glob.glob(os.path.join(SCENES_DIR, 'ep-*.md')))
    for i, path in enumerate(files):
        if i > 0:
            doc.add_page_break()
        render_file(doc, path)

    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Episodes compiled: {len(files)}")

if __name__ == '__main__':
    main()
